import json
import subprocess
import platform
import tempfile
import base64
import shutil
from pathlib import Path
from io import BytesIO

from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SYSTEM_PROMPT = r"""You are a resume data extractor for a staffing agency. You receive raw text from a candidate's resume. Your job is to extract ALL content and structure it into JSON that will be used to generate a cleanly formatted Word document.

RULES — non-negotiable:
1. Extract EVERYTHING from the resume. Do not skip, summarize, condense, or omit any content.
2. Preserve the candidate's exact words in all bullets, descriptions, and text.
3. Fix spelling errors, obvious typos, and grammar mistakes only. Do not rewrite or rephrase anything.
4. Do not invent or add any information that is not in the source text.
5. Standardize date formats: M/YYYY-M/YYYY for ranges (e.g., 6/2013-9/2016), YYYY for single years (e.g., 2015).
6. If information is missing (no email, no phone, no address), set the field to null.
7. Never alter proper nouns, company names, school names, certifications, tools, technologies, or acronyms.
8. Ensure every bullet point ends with a period. If a bullet is missing a period, add one. This is the only content addition allowed.

SECTION IDENTIFICATION:
- Identify all sections in the resume and categorize them.
- Map section names to standard headings when the meaning is obvious:
  "Work Experience" / "Employment History" / "Employment" → "PROFESSIONAL EXPERIENCE"
  "Academic Background" / "Education Background" → "EDUCATION"
  "Technical Skills" / "Core Competencies" → "SKILLS"
  "Certifications & Licenses" → "CERTIFICATIONS"
  "Activities" / "Hobbies" → "INTERESTS"
- If a section does not map cleanly to a standard heading, keep the original heading converted to ALL CAPS.
- Assign a type to each section: education, experience, skills, certifications, interests, volunteer, languages, training, summary, other.
- IMPORTANT: Keep the sections in the same order as they appear in the original resume. Do not reorder them.

ENTRY PARSING (for experience and education sections):
- For each entry, identify: organization name, location (city, state/country), dates, and job title(s)/degree(s).
- One organization can have MULTIPLE roles or titles. Group all roles under a single entry with a "roles" array.
- Each role has a "title" and a "bullets" array.
- Items that appear as bullet points (any bullet character: •, ●, ❖, -, *, ▪, ►, ◆, or similar) go into the "bullets" array as plain text.
- Non-bulleted descriptive text that appears under an entry or role (like a company description or role summary) goes into that role's "description" field.
- Lines like GPA, Dean's List, awards, or honors that are not clearly bullets should still go into "bullets" — they will be rendered as bullet points.

CRITICAL — BULLET SEPARATION RULES (do NOT violate these):
- NEVER combine multiple responsibilities or achievements into a single bullet string. Each distinct task, responsibility, or achievement MUST be its own separate item in the "bullets" array.
- A single bullet should be 1-2 sentences maximum. If you have more than 2 sentences in one bullet, you MUST split it into multiple bullets.
- If the source text has a paragraph describing multiple activities separated by periods, split each sentence into its own bullet.
- The "description" field should ONLY contain a brief introductory line or company description. It should NEVER contain a long paragraph with multiple responsibilities. If you are tempted to put multiple sentences of work responsibilities into "description", put them into "bullets" instead — one bullet per distinct activity.
- When in doubt, create MORE bullets with shorter text rather than FEWER bullets with longer text. Short, scannable bullets are always better than walls of text.

SIMPLE SECTIONS (skills, languages, interests, etc.):
- If a section is a list of items WITHOUT organization/dates structure, use the "items" format.
- Split comma-separated, semicolon-separated, or dash-separated lists into individual items.
- Each item should be a clean, trimmed string.

CONTACT INFORMATION:
- Search the ENTIRE document for contact info — it can appear at the top, bottom, in headers, footers, or scattered throughout.
- Extract: full name, address (or city/state), phone number, email address, LinkedIn URL, personal website.
- The candidate's name is typically the most prominent text, usually at the very top.
- If you find partial info (e.g., city but no full address), include what you find.

WARNINGS:
- Add a warning string for each piece of missing critical info (e.g., "No email address found").
- Add a warning if any dates are ambiguous or missing.
- Add a warning if you made a judgment call on ambiguous content.

OUTPUT FORMAT:
Return ONLY valid JSON. No markdown formatting, no code blocks, no explanation text. Just the raw JSON object.

JSON SCHEMA:
{
  "name": "string — candidate full name",
  "contact": {
    "address": "string or null",
    "phone": "string or null",
    "email": "string or null",
    "linkedin": "string or null",
    "website": "string or null"
  },
  "warnings": ["array of warning strings"],
  "sections": [
    {
      "heading": "SECTION HEADING IN ALL CAPS",
      "type": "education|experience|skills|certifications|interests|volunteer|languages|training|summary|other",
      "entries": [
        {
          "organization": "string — company or school name",
          "organization_details": "string or null — e.g., '(4-year evening program)'",
          "location": "string or null — City, ST or City, Country",
          "dates": "string or null — formatted date or date range",
          "roles": [
            {
              "title": "string or null — job title or degree",
              "description": "string or null — non-bulleted descriptive text",
              "bullets": ["array of bullet point strings"]
            }
          ]
        }
      ],
      "items": ["array of strings — only for simple list sections like skills/languages"]
    }
  ]
}

IMPORTANT NOTES ON THE SCHEMA:
- A section will have EITHER "entries" OR "items", never both. Use "entries" for structured sections (experience, education) and "items" for simple list sections (skills, languages, interests).
- If a section has entries, set "items" to null or omit it.
- If a section has items, set "entries" to null or omit it.
- The "roles" array allows multiple job titles under one organization. Even if there is only one role, still use the array format with one element.
- "description" in a role is for non-bulleted text like company descriptions ("A leading trucking company with 50 employees..."). Preserve this text exactly.
"""


def _get_libreoffice_cmd():
    if platform.system() == "Windows":
        return r"C:\Program Files\LibreOffice\program\soffice.exe"
    return "libreoffice"


def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
    suffix = Path(filename).suffix.lower()
    with tempfile.TemporaryDirectory() as tmpdir:
        tmppath = Path(tmpdir) / filename
        tmppath.write_bytes(file_bytes)

        if suffix == ".docx":
            return _extract_docx(tmppath)
        elif suffix == ".doc":
            docx_path = _convert_doc_to_docx(tmppath)
            return _extract_docx(docx_path)
        elif suffix == ".pdf":
            return _extract_pdf(tmppath)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")


def _extract_docx(path: Path) -> str:
    doc = Document(str(path))
    lines = []
    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            from docx.text.paragraph import Paragraph
            para = Paragraph(child, doc)
            text = para.text.strip()
            if text:
                lines.append(text)
        elif tag == "tbl":
            from docx.table import Table
            table = Table(child, doc)
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append("\t".join(cells))
    return "\n".join(lines)


def _convert_doc_to_docx(path: Path) -> Path:
    cmd = _get_libreoffice_cmd()
    outdir = path.parent
    subprocess.run(
        [cmd, "--headless", "--norestore", "--convert-to", "docx", "--outdir", str(outdir), str(path)],
        check=True,
        capture_output=True,
    )
    result = outdir / f"{path.stem}.docx"
    if not result.exists():
        raise FileNotFoundError(f"Conversion failed: {result} not found")
    return result


def _extract_pdf(path: Path) -> str:
    import pdfplumber
    pages = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n".join(pages)


COMPACT_ADDENDUM = """

ADDITIONAL INSTRUCTION — COMPACT MODE IS ON:
The goal is to produce a resume that fits on ONE PAGE. To achieve this:

1. SHORTEN wordy bullet points. Rewrite them to be concise and professional. Remove filler words, redundant phrases, and unnecessary detail. Keep the core meaning and keywords intact.
2. If a role has a long description paragraph followed by bullets, merge the most important points into fewer, tighter bullets. Drop the least impactful content if needed to hit one page.
3. Remove company description paragraphs entirely (lines like "A leading trucking company with 50 employees..."). The hiring manager does not care about the company description.
4. For simple list sections (skills, software, training), keep them short. Remove obvious filler items (e.g., "Microsoft Office, Internet" — everyone knows these).
5. If there are many roles spanning 20+ years, older roles (10+ years ago) can be reduced to one line: organization, title, dates. No bullets needed for ancient history.
6. NEVER use dashes (- or —) anywhere in the output. Not in bullets, not in descriptions, not in titles. Use commas, periods, or semicolons instead.
7. NEVER invent or add information. You can only remove or shorten existing content.
8. The tone should be professional but not stiff. Write like a competent human, not a corporate robot.
9. Set "compact_mode" to true in the output JSON (add this field at the top level).
"""


def call_llm(resume_text: str, api_key: str, model: str = "x-ai/grok-4.3", compact: bool = False) -> dict:
    client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=api_key)

    prompt = SYSTEM_PROMPT
    if compact:
        prompt += COMPACT_ADDENDUM

    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": prompt},
            {"role": "user", "content": f"Extract and structure this resume:\n\n{resume_text}"},
        ],
        extra_body={"reasoning": {"effort": "high"}},
    )

    content = response.choices[0].message.content
    content = content.strip()
    if content.startswith("```json"):
        content = content[7:]
    if content.startswith("```"):
        content = content[3:]
    if content.endswith("```"):
        content = content[:-3]
    content = content.strip()

    return json.loads(content)


def _set_run_font(run, size_pt=10, bold=False, italic=False, underline=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = r.makeelement(qn("w:rPr"), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")


def _add_tab_stop_right(paragraph, position_inches=7.5):
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = paragraph._element.makeelement(qn("w:pPr"), {})
        paragraph._element.insert(0, pPr)
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = pPr.makeelement(qn("w:tabs"), {})
        pPr.append(tabs)
    tab = tabs.makeelement(
        qn("w:tab"),
        {qn("w:val"): "right", qn("w:pos"): str(int(position_inches * 1440))},
    )
    tabs.append(tab)


def _set_spacing(paragraph, before_pt=0, after_pt=0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before_pt)
    fmt.space_after = Pt(after_pt)


def generate_docx_bytes(data: dict) -> bytes:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    section = doc.sections[0]
    section.top_margin = Inches(0.3125)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    name = data.get("name", "UNKNOWN")
    contact = data.get("contact", {})

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_spacing(name_para, 0, 0)
    _set_run_font(name_para.add_run(name.upper()), size_pt=11, bold=True)

    contact_parts = [contact.get(f) for f in ["address", "phone", "email", "linkedin", "website"] if contact.get(f)]
    if contact_parts:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_spacing(cp, 0, 2)
        _set_run_font(cp.add_run(" ● ".join(contact_parts)), size_pt=10)

    for sec in data.get("sections", []):
        heading_text = sec.get("heading", "OTHER")

        hp = doc.add_paragraph()
        _set_spacing(hp, 4, 1)
        _set_run_font(hp.add_run(heading_text.upper()), size_pt=10, bold=True, underline=True)

        items = sec.get("items")
        if items:
            avg_len = sum(len(i) for i in items) / len(items) if items else 0
            if avg_len < 50:
                ip = doc.add_paragraph()
                _set_spacing(ip, 0, 0)
                clean = [i.rstrip(".").strip() for i in items]
                _set_run_font(ip.add_run(", ".join(clean) + "."), size_pt=10)
            else:
                for item in items:
                    bp = doc.add_paragraph()
                    _set_spacing(bp, 0, 0)
                    bp.paragraph_format.left_indent = Inches(0.5)
                    bp.paragraph_format.first_line_indent = Inches(-0.25)
                    _set_run_font(bp.add_run(f"•  {item}"), size_pt=10)
            continue

        for entry in sec.get("entries", []):
            org = entry.get("organization", "") or ""
            org_details = entry.get("organization_details", "")
            location = entry.get("location", "")
            dates = entry.get("dates", "")

            op = doc.add_paragraph()
            _set_spacing(op, 2, 0)
            _add_tab_stop_right(op, 7.5)

            org_text = org
            if location:
                org_text += f", {location}" if org_text else location
            if org_text:
                _set_run_font(op.add_run(org_text), size_pt=10, bold=True)
            if dates:
                _set_run_font(op.add_run("\t"), size_pt=10)
                _set_run_font(op.add_run(dates), size_pt=10, bold=True)

            if org_details:
                dp = doc.add_paragraph()
                _set_spacing(dp, 0, 0)
                _set_run_font(dp.add_run(org_details), size_pt=10, italic=True)

            for role in entry.get("roles", []):
                title = role.get("title", "")
                description = role.get("description", "")
                bullets = role.get("bullets", [])

                if title:
                    tp = doc.add_paragraph()
                    _set_spacing(tp, 0, 0)
                    _set_run_font(tp.add_run(title), size_pt=10, bold=True)

                all_bullets = []
                if description:
                    all_bullets.append(description)
                all_bullets.extend(bullets)

                for bt in all_bullets:
                    bp = doc.add_paragraph()
                    _set_spacing(bp, 0, 0)
                    bp.paragraph_format.left_indent = Inches(0.5)
                    bp.paragraph_format.first_line_indent = Inches(-0.25)
                    _set_run_font(bp.add_run(f"•  {bt}"), size_pt=10)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def docx_bytes_to_pdf_bytes(docx_bytes: bytes) -> bytes:
    cmd = _get_libreoffice_cmd()
    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = Path(tmpdir) / "resume.docx"
        docx_path.write_bytes(docx_bytes)
        subprocess.run(
            [cmd, "--headless", "--norestore", "--convert-to", "pdf", "--outdir", tmpdir, str(docx_path)],
            check=True,
            capture_output=True,
        )
        pdf_path = Path(tmpdir) / "resume.pdf"
        if not pdf_path.exists():
            raise FileNotFoundError("PDF conversion failed")
        return pdf_path.read_bytes()


def count_pdf_pages(pdf_bytes: bytes) -> int:
    import pdfplumber
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        f.write(pdf_bytes)
        f.flush()
        with pdfplumber.open(f.name) as pdf:
            return len(pdf.pages)


def generate_html_preview(data: dict) -> str:
    name = data.get("name", "UNKNOWN")
    contact = data.get("contact", {})

    html = ['<div style="font-family: \'Times New Roman\', Times, serif; font-size: 10pt; max-width: 7.5in; margin: 0 auto; padding: 20px; background: white; color: black; line-height: 1.3;">']

    html.append(f'<div style="text-align: center; font-size: 11pt; font-weight: bold;">{name.upper()}</div>')

    contact_parts = [contact.get(f) for f in ["address", "phone", "email", "linkedin", "website"] if contact.get(f)]
    if contact_parts:
        html.append(f'<div style="text-align: center; font-size: 10pt; margin-bottom: 4px;">{" ● ".join(contact_parts)}</div>')

    for sec in data.get("sections", []):
        heading = sec.get("heading", "OTHER").upper()
        html.append(f'<div style="font-weight: bold; text-decoration: underline; margin-top: 8px; margin-bottom: 2px;">{heading}</div>')

        items = sec.get("items")
        if items:
            avg_len = sum(len(i) for i in items) / len(items) if items else 0
            if avg_len < 50:
                clean = [i.rstrip(".").strip() for i in items]
                html.append(f'<div>{", ".join(clean)}.</div>')
            else:
                for item in items:
                    html.append(f'<div style="margin-left: 30px; text-indent: -15px;">•&nbsp;&nbsp;{item}</div>')
            continue

        for entry in sec.get("entries", []):
            org = entry.get("organization", "") or ""
            org_details = entry.get("organization_details", "")
            location = entry.get("location", "")
            dates = entry.get("dates", "")

            org_text = org
            if location:
                org_text += f", {location}" if org_text else location

            html.append('<div style="display: flex; justify-content: space-between; margin-top: 4px;">')
            html.append(f'<span style="font-weight: bold;">{org_text}</span>')
            if dates:
                html.append(f'<span style="font-weight: bold; white-space: nowrap;">{dates}</span>')
            html.append('</div>')

            if org_details:
                html.append(f'<div style="font-style: italic;">{org_details}</div>')

            for role in entry.get("roles", []):
                title = role.get("title", "")
                description = role.get("description", "")
                bullets = role.get("bullets", [])

                if title:
                    html.append(f'<div style="font-weight: bold;">{title}</div>')

                all_bullets = []
                if description:
                    all_bullets.append(description)
                all_bullets.extend(bullets)

                for bt in all_bullets:
                    html.append(f'<div style="margin-left: 30px; text-indent: -15px;">•&nbsp;&nbsp;{bt}</div>')

    html.append('</div>')
    return "\n".join(html)
