import sys
import os
import json
import subprocess
import re
from pathlib import Path

from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn

OPENROUTER_KEY = os.environ.get("OPENROUTER_KEY", "")
MODEL = "x-ai/grok-4.3"

SYSTEM_PROMPT = r"""You are a resume data extractor for a staffing agency. You receive raw text from a candidate's resume. Your job is to extract ALL content and structure it into JSON that will be used to generate a cleanly formatted Word document.

RULES — non-negotiable:
1. Extract EVERYTHING from the resume. Do not skip, summarize, condense, or omit any content.
2. Preserve the candidate's exact words in all bullets, descriptions, and text.
3. Fix spelling errors, obvious typos, and grammar mistakes only. Do not rewrite or rephrase anything.
4. Do not invent or add any information that is not in the source text.
5. Standardize date formats: M/YYYY-M/YYYY for ranges (e.g., 6/2013-9/2016), YYYY for single years (e.g., 2015).
6. If information is missing (no email, no phone, no address), set the field to null.
7. Never alter proper nouns, company names, school names, certifications, tools, technologies, or acronyms.
8. Ensure every bullet point ends with a period. If a bullet is missing a period, add one. This is the only content addition allowed.

SECTION IDENTIFICATION:
- Identify all sections in the resume and categorize them.
- Map section names to standard headings when the meaning is obvious:
  "Work Experience" / "Employment History" / "Employment" → "PROFESSIONAL EXPERIENCE"
  "Academic Background" / "Education Background" → "EDUCATION"
  "Technical Skills" / "Core Competencies" → "SKILLS"
  "Certifications & Licenses" → "CERTIFICATIONS"
  "Activities" / "Hobbies" → "INTERESTS"
- If a section does not map cleanly to a standard heading, keep the original heading converted to ALL CAPS.
- Assign a type to each section: education, experience, skills, certifications, interests, volunteer, languages, training, summary, other.
- IMPORTANT: Keep the sections in the same order as they appear in the original resume. Do not reorder them.

ENTRY PARSING (for experience and education sections):
- For each entry, identify: organization name, location (city, state/country), dates, and job title(s)/degree(s).
- One organization can have MULTIPLE roles or titles. Group all roles under a single entry with a "roles" array.
- Each role has a "title" and a "bullets" array.
- Items that appear as bullet points (any bullet character: •, ●, ❖, -, *, ▪, ►, ◆, or similar) go into the "bullets" array as plain text.
- Non-bulleted descriptive text that appears under an entry or role (like a company description or role summary) goes into that role's "description" field.
- Lines like GPA, Dean's List, awards, or honors that are not clearly bullets should still go into "bullets" — they will be rendered as bullet points.

SIMPLE SECTIONS (skills, languages, interests, etc.):
- If a section is a list of items WITHOUT organization/dates structure, use the "items" format.
- Split comma-separated, semicolon-separated, or dash-separated lists into individual items.
- Each item should be a clean, trimmed string.

CONTACT INFORMATION:
- Search the ENTIRE document for contact info — it can appear at the top, bottom, in headers, footers, or scattered throughout.
- Extract: full name, address (or city/state), phone number, email address, LinkedIn URL, personal website.
- The candidate's name is typically the most prominent text, usually at the very top.
- If you find partial info (e.g., city but no full address), include what you find.

WARNINGS:
- Add a warning string for each piece of missing critical info (e.g., "No email address found").
- Add a warning if any dates are ambiguous or missing.
- Add a warning if you made a judgment call on ambiguous content.

OUTPUT FORMAT:
Return ONLY valid JSON. No markdown formatting, no code blocks, no explanation text. Just the raw JSON object.

JSON SCHEMA:
{
  "name": "string — candidate full name",
  "contact": {
    "address": "string or null",
    "phone": "string or null",
    "email": "string or null",
    "linkedin": "string or null",
    "website": "string or null"
  },
  "warnings": ["array of warning strings"],
  "sections": [
    {
      "heading": "SECTION HEADING IN ALL CAPS",
      "type": "education|experience|skills|certifications|interests|volunteer|languages|training|summary|other",
      "entries": [
        {
          "organization": "string — company or school name",
          "organization_details": "string or null — e.g., '(4-year evening program)'",
          "location": "string or null — City, ST or City, Country",
          "dates": "string or null — formatted date or date range",
          "roles": [
            {
              "title": "string or null — job title or degree",
              "description": "string or null — non-bulleted descriptive text",
              "bullets": ["array of bullet point strings"]
            }
          ]
        }
      ],
      "items": ["array of strings — only for simple list sections like skills/languages"]
    }
  ]
}

IMPORTANT NOTES ON THE SCHEMA:
- A section will have EITHER "entries" OR "items", never both. Use "entries" for structured sections (experience, education) and "items" for simple list sections (skills, languages, interests).
- If a section has entries, set "items" to null or omit it.
- If a section has items, set "entries" to null or omit it.
- The "roles" array allows multiple job titles under one organization. Even if there is only one role, still use the array format with one element.
- "description" in a role is for non-bulleted text like company descriptions ("A leading trucking company with 50 employees..."). Preserve this text exactly.
"""


def extract_text(file_path: str) -> str:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".docx":
        return extract_docx(path)
    elif suffix == ".doc":
        docx_path = convert_doc_to_docx(path)
        return extract_docx(docx_path)
    elif suffix == ".pdf":
        return extract_pdf(path)
    else:
        raise ValueError(f"Unsupported file format: {suffix}")


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    lines = []
    body = doc.element.body
    for child in body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            from docx.text.paragraph import Paragraph
            para = Paragraph(child, doc)
            text = para.text.strip()
            if text:
                lines.append(text)
        elif tag == "tbl":
            from docx.table import Table
            table = Table(child, doc)
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append("\t".join(cells))
    return "\n".join(lines)


def convert_doc_to_docx(path: Path) -> Path:
    outdir = path.parent
    subprocess.run(
        [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(outdir),
            str(path),
        ],
        check=True,
        capture_output=True,
    )
    result = outdir / f"{path.stem}.docx"
    if not result.exists():
        raise FileNotFoundError(f"Conversion failed: {result} not found")
    return result


def extract_pdf(path: Path) -> str:
    import pdfplumber

    pages = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n".join(pages)


def call_llm(resume_text: str) -> dict:
    client = OpenAI(base_url="https://openrouter.ai/api/v1", api_key=OPENROUTER_KEY)

    print("Sending to Grok 4.3 (with reasoning)... this may take a minute.")
    response = client.chat.completions.create(
        model=MODEL,
        messages=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": f"Extract and structure this resume:\n\n{resume_text}"},
        ],
        extra_body={"reasoning": {"effort": "high"}},
    )

    content = response.choices[0].message.content

    reasoning = getattr(response.choices[0].message, "reasoning", None)
    if reasoning:
        print(f"Model reasoning length: {len(reasoning)} chars")

    content = content.strip()
    if content.startswith("```json"):
        content = content[7:]
    if content.startswith("```"):
        content = content[3:]
    if content.endswith("```"):
        content = content[:-3]
    content = content.strip()

    try:
        return json.loads(content)
    except json.JSONDecodeError as e:
        print(f"JSON parse error: {e}")
        print(f"Raw response:\n{content[:2000]}")
        raise


def set_run_font(run, size_pt=10, bold=False, italic=False, underline=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = r.makeelement(qn("w:rPr"), {})
        r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")


def add_tab_stop_right(paragraph, position_inches=7.5):
    pPr = paragraph._element.find(qn("w:pPr"))
    if pPr is None:
        pPr = paragraph._element.makeelement(qn("w:pPr"), {})
        paragraph._element.insert(0, pPr)
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = pPr.makeelement(qn("w:tabs"), {})
        pPr.append(tabs)
    tab = tabs.makeelement(
        qn("w:tab"),
        {
            qn("w:val"): "right",
            qn("w:pos"): str(int(position_inches * 1440)),
        },
    )
    tabs.append(tab)


def set_paragraph_spacing(paragraph, before_pt=0, after_pt=0, line_spacing_pt=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before_pt)
    fmt.space_after = Pt(after_pt)
    if line_spacing_pt:
        fmt.line_spacing = Pt(line_spacing_pt)


def generate_docx(data: dict, output_path: str):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.line_spacing = 1.0

    section = doc.sections[0]
    section.top_margin = Inches(0.3125)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    name = data.get("name", "UNKNOWN")
    contact = data.get("contact", {})

    # --- Name line ---
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(name_para, before_pt=0, after_pt=0)
    name_run = name_para.add_run(name.upper())
    set_run_font(name_run, size_pt=11, bold=True)

    # --- Contact line ---
    contact_parts = []
    for field in ["address", "phone", "email", "linkedin", "website"]:
        val = contact.get(field)
        if val:
            contact_parts.append(val)

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(contact_para, before_pt=0, after_pt=2)
        contact_text = " ● ".join(contact_parts)
        contact_run = contact_para.add_run(contact_text)
        set_run_font(contact_run, size_pt=10)

    # --- Sections ---
    for sec in data.get("sections", []):
        heading_text = sec.get("heading", "OTHER")
        sec_type = sec.get("type", "other")

        # Section heading
        heading_para = doc.add_paragraph()
        set_paragraph_spacing(heading_para, before_pt=4, after_pt=1)
        heading_run = heading_para.add_run(heading_text.upper())
        set_run_font(heading_run, size_pt=10, bold=True, underline=True)

        # Simple list sections (skills, languages, interests, etc.)
        items = sec.get("items")
        if items:
            avg_len = sum(len(item) for item in items) / len(items) if items else 0
            if avg_len < 50:
                inline_para = doc.add_paragraph()
                set_paragraph_spacing(inline_para, before_pt=0, after_pt=0)
                clean_items = [item.rstrip(".").strip() for item in items]
                inline_text = ", ".join(clean_items) + "."
                inline_run = inline_para.add_run(inline_text)
                set_run_font(inline_run, size_pt=10)
            else:
                for item in items:
                    bullet_para = doc.add_paragraph()
                    set_paragraph_spacing(bullet_para, before_pt=0, after_pt=0)
                    bullet_para.paragraph_format.left_indent = Inches(0.5)
                    bullet_para.paragraph_format.first_line_indent = Inches(-0.25)
                    bullet_run = bullet_para.add_run(f"•  {item}")
                    set_run_font(bullet_run, size_pt=10)
            continue

        # Structured entries (experience, education)
        entries = sec.get("entries", [])
        for entry in entries:
            org = entry.get("organization", "")
            org_details = entry.get("organization_details", "")
            location = entry.get("location", "")
            dates = entry.get("dates", "")

            # Organization line: "Org, Location\tDates"
            org_para = doc.add_paragraph()
            set_paragraph_spacing(org_para, before_pt=2, after_pt=0)
            add_tab_stop_right(org_para, 7.5)

            org_text = org if org else ""
            if location:
                org_text += f", {location}" if org_text else location

            if org_text:
                org_run = org_para.add_run(org_text)
                set_run_font(org_run, size_pt=10, bold=True)

            if dates:
                tab_run = org_para.add_run("\t")
                set_run_font(tab_run, size_pt=10)
                date_run = org_para.add_run(dates)
                set_run_font(date_run, size_pt=10, bold=True)

            if org_details:
                details_para = doc.add_paragraph()
                set_paragraph_spacing(details_para, before_pt=0, after_pt=0)
                details_run = details_para.add_run(org_details)
                set_run_font(details_run, size_pt=10, italic=True)

            # Roles
            roles = entry.get("roles", [])
            for role in roles:
                title = role.get("title", "")
                description = role.get("description", "")
                bullets = role.get("bullets", [])

                if title:
                    title_para = doc.add_paragraph()
                    set_paragraph_spacing(title_para, before_pt=0, after_pt=0)
                    title_run = title_para.add_run(title)
                    set_run_font(title_run, size_pt=10, italic=False)

                all_bullets = []
                if description:
                    all_bullets.append(description)
                all_bullets.extend(bullets)

                for bullet_text in all_bullets:
                    bullet_para = doc.add_paragraph()
                    set_paragraph_spacing(bullet_para, before_pt=0, after_pt=0)
                    bullet_para.paragraph_format.left_indent = Inches(0.5)
                    bullet_para.paragraph_format.first_line_indent = Inches(-0.25)
                    bullet_run = bullet_para.add_run(f"•  {bullet_text}")
                    set_run_font(bullet_run, size_pt=10)

    # --- Warnings ---
    warnings = data.get("warnings", [])
    if warnings:
        print("\nWarnings from extraction:")
        for w in warnings:
            print(f"  - {w}")

    doc.save(output_path)
    print(f"\nFormatted resume saved: {output_path}")


def main():
    if len(sys.argv) < 2:
        print("Usage: python poc.py <resume_file> [output_file]")
        print("Supported formats: .doc, .docx, .pdf")
        sys.exit(1)

    input_file = sys.argv[1]
    if not Path(input_file).exists():
        print(f"File not found: {input_file}")
        sys.exit(1)

    output_file = sys.argv[2] if len(sys.argv) > 2 else str(
        Path(input_file).parent / f"{Path(input_file).stem}_formatted.docx"
    )

    print(f"Input:  {input_file}")
    print(f"Output: {output_file}")
    print()

    print("Step 1: Extracting text...")
    text = extract_text(input_file)
    print(f"Extracted {len(text)} characters.")
    print()

    print("Step 2: Calling LLM for extraction + structuring...")
    data = call_llm(text)

    json_debug = Path(input_file).parent / f"{Path(input_file).stem}_debug.json"
    with open(json_debug, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    print(f"Debug JSON saved: {json_debug}")
    print()

    print("Step 3: Generating formatted DOCX...")
    generate_docx(data, output_file)

    print("\nDone!")


if __name__ == "__main__":
    main()
